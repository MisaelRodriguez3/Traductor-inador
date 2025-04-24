from docx.document import Document
from docx.shared import Pt, RGBColor
from app.utils.error_handler import handle_error

def add_watermark(doc: Document, text: str) -> None:
    """Adds a styled watermark to the footer of a Word document.
    
    Applies consistent watermark styling to the first section's footer:
    - Century Gothic font (10pt)
    - Custom gold color (RGB: 239, 210, 132)
    - Bold formatting
    - Left-aligned text

    Args:
        doc (Document): python-docx Document object to modify
        text (str): Watermark text content to display

    Raises:
        Exception: Propagates any errors during processing to handler
        ValueError: If document contains no sections
        AttributeError: If font configuration fails

    Examples:
        >>> from docx import Document
        >>> doc = Document()
        >>> add_watermark(doc, "Confidential")
        >>> doc.save("document_with_watermark.docx")

    Notes:
        - Modifies document in-place
        - Requires Century Gothic font installed for exact appearance
        - Uses first section only for watermark placement
        - Overwrites existing content in footer's first paragraph
    """
    try: 
        if not doc.sections:
            raise ValueError("Document contains no sections for watermarking")

        section = doc.sections[0]
        footer = section.footer

        # Get or create first footer paragraph
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT

        # Configure watermark styling
        run = paragraph.add_run(text)
        run.font.name = 'Century Gothic'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(239, 210, 132)  # Gold color
        run.bold = True

        # Force style element creation
        run._element.get_or_add_rPr()
        
    except Exception as e:
        handle_error(e)