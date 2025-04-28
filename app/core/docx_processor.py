import os
from typing import Callable, Set
from docx import Document
from app.exceptions.document import (
    DocumentNotFound,
    DocumentReadError,
    DocumentWriteError,
    ParagraphTranslationError,
)

class DocxProcessor:
    """Processes DOCX documents for translation with progress tracking and resume capabilities.   
    
    Features:
    - Preserves original document formatting and structure
    - Checkpoint system for resuming interrupted translations
    - Page skipping functionality
    - Chunked translation to handle API limits
    
    Attributes:
        translator (TranslationService): Translation service instance with translate() method
        chunk_size (int): Maximum characters per translation chunk (default: 200)
    
    Raises:
        DocumentNotFound: When input file is not found
        DocumentReadError: On document loading failures
        DocumentWriteError: On document saving failures
        ParagraphTranslationError: When paragraph translation fails
    """

    def __init__(self, translator: object, chunk_size: int = 200):
        """Initializes the document processor with translation service and configuration.
        
        Args:
            translator: Translation service implementing translate(text, src_lang, dest_lang)
            chunk_size: Maximum character count per translation chunk (default: 200)
        """
        self.translator = translator
        self.chunk_size = chunk_size

    def process_document(
        self,
        input_path: str,
        output_path: str,
        lang_from: str,
        lang_to: str,
        progress_callback: Callable[[int, int], None] | None = None,
        skip_pages: Set[int] = set()
    ) -> None:
        """Main method to process and translate a DOCX document.
        
        Args:
            input_path (str): Path to source DOCX file
            output_path (str): Path for translated DOCX file
            lang_from (str): Source language code (ISO 639-1)
            lang_to (str): Target language code (ISO 639-1)
            progress_callback (Callable[[int, int], None], optional): Optional callback for progress updates (processed, total)
            skip_pages (set[int], optional): Set of page numbers to skip in translation
        
        Raises:
            DocumentNotFound: If input file doesn't exist
            DocumentReadError: If document can't be loaded
            DocumentWriteError: If document can't be saved
            ParagraphTranslationError: If any paragraph fails to translate
        """
        checkpoint_path = f"{output_path}.checkpoint"
        start_index, current_page = self._load_checkpoint(checkpoint_path)
        
        try:
            doc = self._load_document(input_path)
            paragraphs = self._extract_all_paragraphs(doc)
            total = len(paragraphs)
            
            for idx in range(start_index, total):
                paragraph = paragraphs[idx]
                current_page = self._handle_page_breaks(paragraph, current_page)
                
                if current_page in skip_pages:
                    continue
                
                try:
                    self._translate_paragraph(paragraph, lang_from, lang_to)
                except Exception as e:
                    self._save_progress(doc, output_path, checkpoint_path, idx)
                    raise ParagraphTranslationError(f"Paragraph {idx+1} error: {e}")
                
                self._update_checkpoint(checkpoint_path, idx + 1)
                self._report_progress(progress_callback, idx + 1, total)
            
            self._finalize_output(doc, output_path, checkpoint_path)
        
        except Exception as e:
            if isinstance(e, (DocumentNotFound, DocumentReadError, DocumentWriteError)):
                raise
            raise DocumentWriteError(f"Unexpected error: {e}")

    def _load_document(self, path: str):
        """Loads DOCX document from file path.
        
        Args:
            path (str): Path to document file
            
        Returns:
            Document: Loaded document object
            
        Raises:
            DocumentNotFound: If file not found
            DocumentReadError: For other loading errors
        """
        try:
            return Document(path)
        except FileNotFoundError:
            raise DocumentNotFound(f"File not found: {path}")
        except Exception as e:
            raise DocumentReadError(f"Error loading document: {e}")

    def _extract_all_paragraphs(self, doc) -> list:
        """Extracts all paragraphs including tables and their cells.
        
        Args:
            doc (Document): python-docx Document object
            
        Returns:
            list[Paragraph]: All paragraph objects in document
        """
        return (
            list(doc.paragraphs)
            + [
                para
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
                for para in cell.paragraphs
            ]
        )

    def _translate_paragraph(self, paragraph, lang_from: str, lang_to: str) -> None:
        """Translates a paragraph while preserving formatting runs.
        
        Args:
            paragraph (Paragraph): docx Paragraph object to translate
            lang_from (str): Source language code
            lang_to (str): Target language code
        """
        if not paragraph.text.strip():
            return
        
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            
            chunks = self._split_into_chunks(run.text)
            translated = [self.translator.translate(chunk, lang_from, lang_to) for chunk in chunks]
            run.text = "".join(translated)

    def _split_into_chunks(self, text: str) -> list[str]:
        """Splits text into chunks respecting word boundaries and size limit.
        
        Args:
            text (str): Input text to split
            
        Returns:
            list[str]: Text chunks under size limit
        """
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in text.split():
            word_length = len(word) + 1  # Include space
            if current_length + word_length > self.chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = [word]
                current_length = word_length
            else:
                current_chunk.append(word)
                current_length += word_length
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks

    def _handle_page_breaks(self, paragraph, current_page: int) -> int:
        """Tracks page numbers based on paragraph breaks.
        
        Args:
            paragraph (Paragraph): Paragraph to check for page breaks
            current_page (int): Current page number
            
        Returns:
            int: Updated page number
        """
        if self._has_page_break(paragraph):
            return current_page + 1
        return current_page

    def _has_page_break(self, paragraph) -> bool:
        """Detects page breaks in paragraph using XML structure.
        
        Args:
            paragraph (Paragraph): Paragraph to check
            
        Returns:
            bool: True if contains page break
        """
        return any('w:lastRenderedPageBreak' in run._element.xml for run in paragraph.runs)

    def _load_checkpoint(self, checkpoint_path: str) -> tuple[int, int]:
        """Loads translation progress from checkpoint file.
        
        Args:
            checkpoint_path (str): Path to checkpoint file
            
        Returns:
            tuple[int, int]: (start_index, current_page)
        """
        if os.path.exists(checkpoint_path):
            try:
                with open(checkpoint_path, "r") as f:
                    return int(f.read().strip()), 1  # Page tracking resets
            except (ValueError, IOError):
                return 0, 1
        return 0, 1

    def _update_checkpoint(self, checkpoint_path: str, index: int) -> None:
        """Updates checkpoint file with current progress index.
        
        Args:
            checkpoint_path (str): Path to checkpoint file
            index (int): Current paragraph index
            
        Raises:
            DocumentWriteError: If checkpoint update fails
        """
        try:
            with open(checkpoint_path, "w") as f:
                f.write(str(index))
        except IOError as e:
            raise DocumentWriteError(f"Checkpoint update failed: {e}")

    def _save_progress(self, doc, output_path: str, checkpoint_path: str, index: int) -> None:
        """Saves current progress and checkpoint during error handling.
        
        Args:
            doc (Document): Document object to save
            output_path (str): Output file path
            checkpoint_path (str): Checkpoint file path
            index (int): Current progress index
            
        Raises:
            DocumentWriteError: If save operation fails
        """
        try:
            doc.save(output_path)
            self._update_checkpoint(checkpoint_path, index)
        except Exception as e:
            raise DocumentWriteError(f"Error saving progress: {e}")

    def _finalize_output(self, doc, output_path: str, checkpoint_path: str) -> None:
        """Saves final document and cleans up checkpoint.
        
        Args:
            doc (Document): Document object to save
            output_path (str): Output file path
            checkpoint_path (str): Checkpoint file path
            
        Raises:
            DocumentWriteError: If final save fails
        """
        try:
            doc.save(output_path)
            if os.path.exists(checkpoint_path):
                os.remove(checkpoint_path)
        except Exception as e:
            raise DocumentWriteError(f"Final save failed: {e}")

    def _report_progress(
        self,
        callback: Callable[[int, int], None] | None,
        processed: int,
        total: int
    ) -> None:
        """Executes progress callback if provided.
        
        Args:
            callback (Callable[[int, int], None] | None): Progress reporting function
            processed (int): Number of processed items
            total (int): Total number of items
        """
        if callback:
            callback(processed, total)